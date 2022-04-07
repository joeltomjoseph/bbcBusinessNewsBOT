import urllib.request as ul
from bs4 import BeautifulSoup as soup
from docx import Document
from docx.shared import Inches
import random
import time
import json

def getInfo():
    url = "https://www.bbc.co.uk/news/business"
    url2 = ""
    req = ul.Request(url, headers={'User-Agent':'Chrome'})
    req2 = ul.Request(url2, headers={'User-Agent':'Chrome'})
    client = ul.urlopen(req)
    client2 = ul.urlopen(req2)
    htmldata = client.read()
    weatherdata = client2.read()
    client.close()
    client2.close()

    htmlsouped = soup(htmldata, "html.parser")
    newslocater = htmlsouped.find('div', {"class":"gel-layout gel-layout--equal"})

    weatherlocater = json.loads(weatherdata)
    #print(weatherlocater)

    allweatherinfo = []
    
    for unixtime, weatherinfo in [[[hourlytime['dt'] for hourlytime in weatherlocater['hourly']], [hourlyweather['weather'] for hourlyweather in weatherlocater['hourly']]]]:
        allweatherinfo.append([unixtime,weatherinfo])
    print(allweatherinfo[0][1])

    titles = []
    for items in newslocater.find_all('h3'):
            title = items.text
            titles.append(title)
            #print(title)
    titles.pop(5)

    summaries = []
    for items in newslocater.find_all('p'):
            summary = items.text
            summaries.append(summary)
            #print(summary)

    images = []
    for items in newslocater.find_all('img', {"class":"lazyloaded"}):### not exactly working
            image = items.get('src')
            images.append(image)
            #print(image)

    return titles, summaries, images

def createDocument():

    titles, summaries, images = getInfo()
    
    document = Document()

    document.add_heading('BBC Business News Report', 0)

    p = document.add_paragraph('Here`s the ')
    p.add_run('news').bold = True
    p.add_run(' for ')
    p.add_run('today ').italic = True
    
    for title in titles:
        document.add_heading(title, level=1)
        document.add_paragraph(summaries[titles.index(title)], style='Intense Quote')

    document.add_picture(f"Documents/Projects/Coding Projects/bbcBusinessNewsBOT/{random.randint(1,4)}.jpg", width=Inches(3))

    document.save('Documents/Projects/Coding Projects/bbcBusinessNewsBOT/newsReport.docx')
getInfo()
#createDocument()